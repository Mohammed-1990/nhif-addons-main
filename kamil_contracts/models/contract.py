# -*- coding: utf-8 -*-

from odoo import models, fields, api
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_DIRECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Mm
from docx.shared import Inches
from odoo.exceptions import ValidationError, Warning

import shutil
import os
import getpass
import platform
import docx 
import os
import io
import re
import base64


class kamilProjectContract(models.Model):
    _name = 'kamil.contracts.contract'
    _description = 'Contract'
    _order = 'id desc'
    _inherit = ['mail.thread','mail.activity.mixin', 'portal.mixin']


    # _name = 'kamil.contracts.project.contract'
    state = fields.Selection([
        ('draft', 'Draft'),
        ('submit','Submit'), 
        ('study', 'Study'),
        ('second_party', 'Second party Sign.'),
        ('first_party','First party sign..'),
        ('Witness_sig','Witness Signature'),
        ('attestation','Attestation'),
        ('del_copy_sec','Delivery a copy to S.P'),
        ('del_copy_req', 'Deliver Copy to Requester'),
        ('done', 'Done'),
        ('open', 'Running'),
        ('pending', 'To Renew'),
        ('close', 'Expired'),
        ('terminated','Terminated'),
        ('cancel', 'Canceld'), ],"state", track_visibility="always", default="draft")

    def cleanhtml(self,raw_html):
        # cleanr = re.compile('<.*?>')
        cleanr = re.compile(r'<[^>]+>')
        cleantext = re.sub(cleanr, '', raw_html)

        
        
        return cleantext

    def get_contract_template(self):
        # docx_stream = io.BytesIO()
        doc = docx.Document()

        # document = Document()
        section = doc.sections[0]
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.left_margin = Mm(25.4)
        section.right_margin = Mm(25.4)
        section.top_margin = Mm(25.4)
        section.bottom_margin = Mm(25.4)
        section.header_distance = Mm(12.7)
        section.footer_distance = Mm(12.7)

        
        nll= doc.add_paragraph('\n\n\n')
        delay_mulct_html= doc.add_paragraph('')
        doc.add_heading(self.name, 0)
        


        preface= doc.add_paragraph('')
        preface.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pref=self.cleanhtml(self.preface)
        preface.add_run('الديباجة \n').bold = True
        preface.add_run(pref)
        preface.add_run('\n')

        
        preamble= doc.add_paragraph('')
        preamble.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pream=self.cleanhtml(self.preamble)
        preamble.add_run('التمهيد \n').bold = True
        preamble.add_run(pream)
        preamble.add_run('\n')
        
        explain= doc.add_paragraph('')
        explain.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        expl=self.cleanhtml(self.explain)
        explain.add_run('التفسير \n').bold = True
        explain.add_run(expl)
        explain.add_run('\n')

        contract_docx= doc.add_paragraph('')
        contract_docx.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        contract_doc=self.cleanhtml(self.contract_docx)
        contract_docx.add_run('مستندات العقد \n').bold = True
        contract_docx.add_run(contract_doc)
        contract_docx.add_run('\n')
        
        
        first_party_obligations= doc.add_paragraph('')
        first_party_obligations.alignment  = WD_ALIGN_PARAGRAPH.RIGHT
        first_party_obl=self.cleanhtml(self.first_party_obligations)
        first_party_obligations.add_run('إلتزامات الطرف الاول \n').bold = True
        first_party_obligations.add_run(first_party_obl)
        first_party_obligations.add_run('\n')
        

        second_party_obligations= doc.add_paragraph('')
        second_party_obligations.alignment  = WD_ALIGN_PARAGRAPH.RIGHT
        second_party_obl=self.cleanhtml(self.second_party_obligations)
        second_party_obligations.add_run('التزامات الطرف الثاني \n').bold = True
        second_party_obligations.add_run(second_party_obl)
        second_party_obligations.add_run('\n')
        
        general_provisions= doc.add_paragraph('')
        general_provisions.alignment  = WD_ALIGN_PARAGRAPH.RIGHT
        general_provision=self.cleanhtml(self.general_provisions)
        general_provisions.add_run('أحكام عامة \n').bold = True
        general_provisions.add_run(general_provision)
        general_provisions.add_run('\n')
        
        delay_mulct_html= doc.add_paragraph('')
        delay_mulct_html.alignment  = WD_ALIGN_PARAGRAPH.RIGHT
        delay_mulct_=self.cleanhtml(self.delay_mulct_html)
        delay_mulct_html.add_run('المخالفات وغرامات التاخير  \n').bold = True
        delay_mulct_html.add_run(delay_mulct_)
        delay_mulct_html.add_run('\n')
        
        conflict_resolution_html= doc.add_paragraph('')
        conflict_resolution_html.alignment  = WD_ALIGN_PARAGRAPH.RIGHT
        conflict_resolution_=self.cleanhtml(self.conflict_resolution_html)
        conflict_resolution_html.add_run('آلية فض النزاع \n').bold = True
        conflict_resolution_html.add_run(conflict_resolution_)
        conflict_resolution_html.add_run('\n')
                
        parties_table = doc.add_table(rows=3, cols=4)  #Create a table with and set the rows and columns        
        parties_table.direction = WD_TABLE_DIRECTION.RTL
        parties_table.style = "Table Grid"   # Define the table style.  You can set any style defined in the styles files of the library
        hdr_cells = parties_table.rows[0].cells
        
        # hdr_cells[1].paragraphs[1].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
        # hdr_cells[2].paragraphs[2].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
        # hdr_cells[3].paragraphs[3].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
        
        hdr_cells[0].text = self.second_party_ch.name
        hdr_cells[1].text = 'الطرف الثاني'
        hdr_cells[2].text = str(self.ntfs.name)
        hdr_cells[3].text = 'الطرف الاول'

        tbl_cells_2 = parties_table.rows[1].cells
        tbl_cells_2[0].text = self.second_party_name.name
        tbl_cells_2[1].text = 'ممثل الطرف الثاني'
        tbl_cells_2[2].text = str(self.name_ntfs.name)
        tbl_cells_2[3].text = 'ممثل الطرف الاول'


        tbl_cells_3 = parties_table.rows[2].cells
        tbl_cells_3[0].text = self.second_party_obj
        tbl_cells_3[1].text = 'صفة الطرف الثاني'
        tbl_cells_3[2].text = str(self.name_ntfs_obj.name)
        tbl_cells_3[3].text = 'صفة الطرف الاول'

        new_line= doc.add_paragraph('')

        witnesses_table = doc.add_table(rows=3, cols=4)  #Create a table with and set the rows and columns
        witnesses_table.direction = WD_TABLE_DIRECTION.LTR
        witnesses_table.style = "Table Grid"   # Define the table style.  You can set any style defined in the styles files of the library
        hdr_cells = witnesses_table.rows[0].cells
        
        # hdr_cells[1].paragraphs[1].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
        # hdr_cells[2].paragraphs[2].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
        # hdr_cells[3].paragraphs[3].paragraph_format.alignment = WD_TABLE_ALIGNMENT.RIGHT
        
        hdr_cells[0].text = str(self.second_witnes_name)
        hdr_cells[1].text = 'الشاهد الثاني'
        hdr_cells[2].text = str(self.first_witnes_name)
        hdr_cells[3].text = 'الشاهد الاول'

        tbl_cells_2 = witnesses_table.rows[1].cells
        tbl_cells_2[0].text = str(self.second_witnes_id)
        tbl_cells_2[1].text = 'نوع ﻹثبات الشخصية'
        tbl_cells_2[2].text = str(self.first_witnes_id)
        tbl_cells_2[3].text = 'نوع ﻹثبات الشخصية'


        tbl_cells_3 = witnesses_table.rows[2].cells
        tbl_cells_3[0].text = str('')
        tbl_cells_3[1].text = 'تاريخها'
        tbl_cells_3[2].text = str('')
        tbl_cells_3[3].text = 'تاريخها'


        username = getpass.getuser()
        op_system=platform.system()
        if op_system=='Linux':
            os.chdir('/home/'+username+'/Downloads')
            doc.save(str(self.name+'.docx'))
            
            read_bytes=self.get_bytes_from_file((str('/home/'+username+'/Downloads/'+self.name+'.docx')))
            attachment = {'name': self.name,    
            'datas': base64.b64encode(read_bytes), 
            'datas_fname': self.name+'.docx', 
            'res_model': self._name, 
            'res_id': self.id,
            'mimetype': 'application/x-doc'}

            self.env['ir.attachment'].create(attachment)
            os.remove(str('/home/'+username+'/Downloads/'+self.name+'.docx'))
            
            return attachment    
        
        else:
            #username=os.path.split(userhome)[-1]
            #username=os.getusername()
            os.chdir(r'C:/Users/Public/Downloads')
            doc.save(str(self.name+'.docx'))
            
            read_bytes=self.get_bytes_from_file((str('C:/Users/Public/Downloads/'+self.name+'.docx')))
            attachment = {'name': self.name,    
            'datas': base64.b64encode(read_bytes), 
            'datas_fname': self.name+'.docx', 
            'res_model': self._name, 
            'res_id': self.id,
            'mimetype': 'application/x-doc'}

            self.env['ir.attachment'].create(attachment)
            
                
            os.remove(str('C:/Users/Public/Downloads/'+self.name+'.docx'))

            return attachment    
            
        
        

    def get_bytes_from_file(self,filename):  
        return open(filename, "rb").read()    


    
    @api.multi
    def submit(self):
        self.write({'state':'submit'})
    @api.multi
    def study(self):
        self.write({'state':'study'})
    @api.multi
    def second_party(self):
        self.write({'state':'second_party'})
    @api.multi
    def first_party(self):
        self.write({'state':'first_party'})
    
    @api.multi
    def Witness_sig(self):
        self.write({'state':'Witness_sig'})
        
    @api.multi
    def attestation(self):
        self.write({'state':'attestation'})

    @api.multi
    def del_copy_sec(self):
        self.write({'state':'del_copy_sec'})

    @api.multi
    def del_copy_req(self):
        self.write({'state':'del_copy_req'})

    @api.multi
    def done(self):
        #self.employee_id.toggle_active()
        self.write({'state':'done'})

    @api.multi
    def open(self):
        for penalty in self.panalty_ids:
            penalty.write({'opened' : True})
        self.write({'state':'open'})
        if self.contract_type=='Purchase' or self.contract_type=='Service':
            self.purchase_order_id.button_confirm()

    @api.multi
    def pending(self):
        for penalty in self.panalty_ids:
            penalty.write({'opened' : False})
        self.write({'state':'pending'})
    
    @api.multi
    def close(self):
        self.write({'state':'close'})
    @api.multi
    def cancel(self):
        self.write({'state':'cancel'})

    @api.multi
    def terminated(self):
        self.write({'state':'terminated'})
     
    
    # @api.multi
    # @api.constrains('contract_amount')
    # def validat_contract_amount(self):
    #     for rec in self:
    #         if rec.contract_amount == 0:
    #             raise ValidationError(_('Error contract amount cannot be zero'))
    

    @api.multi
    @api.onchange('preface_id')
    def onchnage_preface(self):
        if self.preface_id:
            self.preface = self.preface_id.content_p

    
    @api.multi
    @api.onchange('introduction_id')
    def onchnage_introduction(self):
        if self.introduction_id:
            self.introduction = self.introduction_id.content_introduct
    
    @api.multi
    @api.onchange('first_party_obligation_sel')
    def onchnage_first(self):
        if self.first_party_obligation_sel:
            self.first_party_obligations = self.first_party_obligation_sel.content_first
    
    @api.multi
    @api.onchange('second_party_obligations_sel')
    def onchnage_second(self):
        if self.second_party_obligations_sel:
            self.second_party_obligations = self.second_party_obligations_sel.content_second
    

    @api.multi
    @api.onchange('preamble_id')
    def onchnage_preamble(self):
        if self.preamble_id:
            self.preamble = self.preamble_id.content_introduct
  
    @api.multi
    @api.onchange('explain_id')
    def onchnage_explain(self):
        if self.explain_id:
            self.explain = self.explain_id.content
   
    @api.multi
    @api.onchange('delay_mulct_id')
    def onchnage_delay(self):
        if self.delay_mulct_id:
            self.delay_mulct_html = self.delay_mulct_id.content
    
    @api.multi
    @api.onchange('conflict_resolution_id')
    def onchnage_conflict(self):
        if self.conflict_resolution_id:
            self.conflict_resolution_html = self.conflict_resolution_id.content
    
 
   
    @api.multi
    @api.onchange('contract_docx_id')
    def onchnage_docx(self):
        if self.contract_docx_id:
            self.contract_docx = self.contract_docx_id.content
    
    @api.multi
    @api.onchange('contract_end_date')
    def onchange_contract_start_date(self):
        if self.contract_end_date and self.contract_start_date:

            if self.contract_end_date < self.contract_start_date:
                raise exceptions.ValidationError('Please select a date equal/or greater than the contract Start date')
                

    @api.multi
    def check_contract_end_date(self):
        for rec in self.env['kamil.contracts.contract'].search([('state','=','open')]):
            if rec.contract_end_date:
                if rec.contract_end_date < fields.Date.today():
                    rec.state = 'close'
    @api.multi
    def check_contract_start_date(self):
        # print ("Start contracts cron job \n\n")
        for rec in self.env['kamil.contracts.contract'].search([('state','=','done')]):
        # for rec in self.env['kamil.contracts.contract'].search([]):
            # print (rec.contract_start_date,"\n\n")
            if rec.contract_start_date:
                # print (rec.contract_start_date,"\n\n")
                # print (fields.Date.today(),"\n\n")
                if (rec.contract_start_date == fields.Date.today()):
                    rec.state = 'open'


    @api.multi
    def contract_extensions(self):
        self.ensure_one()
        return{
            'name': ('Extensions'),
            # 'domain':[('contract_id','in',"['6']")],
            'domain': [('contract_id', '=', self.id)],
            'view_type':'form',
            'res_model':'kamil.contracts.extensions',
            'view_id':False,
            'view_mode':'tree,form',
            'type':'ir.actions.act_window',            
            # 'context': {'default_contract_id': self.id, 'create':'0'},
            'context': "{'create': False,'edit': False}"
        }
    
    
   
    

    @api.multi
    def contract_ammedments(self):
        return{
            'name': ('Ammedment'),
            'domain': [('contract_id', '=', self.id)],
            'view_type':'form',
            'res_model':'kamil.contracts.contract',
            'view_id':False,
            # 'view_id':'kamil_contracts.kamil_general_contract_form_view',
            # 'view_mode':'form',
            'view_mode':'tree,form',
            'type':'ir.actions.act_window',
            'context':{'default_contract_type':'Ammendment',
                        'default_preamble':self.preamble,    
                        'default_first_party':self.first_party,
                        'default_first_party_obligations':self.first_party_obligations,
                        'default_second_party_obligations':self.second_party_obligations,
                        'default_contract_docx':self.contract_docx,
                        'default_contract_amount':self.contract_amount,
                        'default_general_provisions':self.general_provisions,
                        'default_delay_mulct_html':self.delay_mulct_html,
                        'default_conflict_resolution_html':self.conflict_resolution_html
                        ,
                        'default_ntfs':self.ntfs.id,
                        'default_name_ntfs':self.name_ntfs.id,
                        'default_name_ntfs_obj':self.name_ntfs_obj.id,
                        'default_second_party_name':self.second_party_name.id,
                        'default_second_party_ch':self.second_party_ch.id,
                        'default_second_party_obj':self.second_party_obj,




            },
        }


    @api.multi
    def contract_products(self):
        return{
            'name': ('Order lines'),
            'domain':[],
            'view_type':'form',
            'res_model':'purchase.order.line',
            'view_id':False,
            'view_mode':'tree,form',
            'type':'ir.actions.act_window',
        }


    @api.model
    def _get_default_nhif(self):
        if self.contract_id:
            return self.contract_id.ntfs.id

    name = fields.Char(string='Contract Name', required=True, track_visibility="always")
    company_id = fields.Many2one('res.company', string='Company', required=True,
        copy=False, default=lambda self: self.env['res.company']._company_default_get())
    
    contract_id= fields.Many2one('kamil.contracts.contract', 'Contract ID', ondelete='cascade', track_visibility="always")
    contract_ids= fields.One2many('kamil.contracts.contract', 'contract_id', 'Contracts')

    contract_type = fields.Char(string='Contract Type',  track_visibility="always", default='Ammendment')
    
    # For medical contracts
    sector = fields.Selection([('m', 'public_sector'), ('f', 'Private')], string='Sector', track_visibility="always")
    specialty = fields.Selection([('n', 'Multidisciplinary'), ('s', 'Specialized'), ('g', 'Specialized/Multidisciplinary')], string='Type', track_visibility="always")

    preamble_id = fields.Many2one('kamil.contracts.article.introduct')
    preamble = fields.Html('', track_visibility="always")
    
    preface_id = fields.Many2one('kamil.contracts.article.preface')
    preface = fields.Html('', track_visibility="always")

    introduction_id = fields.Many2one('kamil.contracts.article.introduct')
    introduction = fields.Html('Introduction', track_visibility="always")    
 
  
    explain = fields.Html('Explaination')
    explain_id = fields.Many2one('kamil.contracts.article.explain')
    
     
    contract_start_date = fields.Date('Contract start Date',required=True ,default=fields.Date.today(), track_visibility="always")
    contract_end_date = fields.Date('Contract end Date', track_visibility="always")

    
    contract_docx_id = fields.Many2one('kamil.contracts.article.documents')
    contract_docx = fields.Html('Contract documents')

    delay_mulct_id = fields.Many2one('kamil.contracts.article.delay_panalty')
    delay_mulct_html = fields.Html('')

    conflict_resolution_html = fields.Html('')
    conflict_resolution_id = fields.Many2one('kamil.contracts.article.conflict_resolution')
 
    contract_duration = fields.Char(' ')
    effective_date = fields.Date('')
  

    period_name = fields.Selection([('m', 'Month/s'), ('y', 'Year/s')], default="m", track_visibility="always")
    period = fields.Selection([('o', '1'), ('s', '2'),('t','3'), ('o4', '4'), ('o5', '5'), ('o6', '6'), ('o7', '7'), ('o8', '8'), ('o9', '9'), ('o10', '10'), ('o11', '11'), ('o12', '12')],string="Period ", track_visibility="always")
    
    # amount_contract = fields.Integer('Amount')
    first_party_obligation_sel = fields.Many2one('kamil.contracts.article.first')
    first_party_obligations = fields.Html('', track_visibility="always")    

    second_party_obligations_sel = fields.Many2one('kamil.contracts.article.second')
    second_party_obligations = fields.Html('', track_visibility="always")
        
    public_judge = fields.Html('', track_visibility="always")
    general_provisions = fields.Html('General Provisions')

    ntfs = fields.Many2one('res.company',string='NHIF', required=True)  
    # , default=lambda self: self._get_default_nhif()
    name_ntfs = fields.Many2one('hr.employee',string='Rep. Name', required=True,)
    name_ntfs_obj = fields.Many2one('hr.function',string='First Rep. Capacity', readonly=True, required=True ,related="name_ntfs.functional_id")

    

    second_party_ch = fields.Many2one('res.partner',string='Name', required=True, track_visibility="always")
    second_party_name = fields.Many2one('res.partner',string='Second Rep. Name', required=True, track_visibility="always")
    
    # second_party_obj = fields.Char(related='second_party_name.title')
    # second_party_obj = fields.Many2one('res.partner',string='Rep. Capacity', required=True, track_visibility="always", related="second_party_name.function", readonly=True)
    second_party_obj = fields.Char(string='Rep. Capacity', track_visibility="always", readonly=True, related="second_party_name.function")
   
    second_party_ch_id= fields.Integer()
    

    contract_amount=fields.Integer('Contract Amount', track_visibility="always")
    contracts_date = fields.Datetime('Contract Date', default=lambda self: fields.Datetime.now(), track_visibility="always")
    contract_period = fields.Integer('Contract Period' , track_visibility="always")
   
    # Rewrite , Verifiy, Approved
    is_contract_active= fields.Boolean('Is Active', default=True, track_visibility="always")

    first_witnes_name = fields.Char('First Witness Name', track_visibility="always")
    first_witnes_signature = fields.Char('First Witness Signature')
    first_witnes_id = fields.Selection([('national_ID','ID Card'),('passport','Passport')], track_visibility="always")


    second_witnes_name = fields.Char('Second Witness Name', track_visibility="always")
    second_witnes_signature = fields.Char('Second Witness Signature')
    second_witnes_id = fields.Selection([('national_ID','ID Card'),('passport','Passport')], track_visibility="always")

    
    # For May mode
    payment_mode_ids = fields.One2many('kamil.contracts_waypay','contract_id','contract_id_pay_rel')
    # payment_mode_ids = fields.Many2many('kamil.contracts_waypay','rel_kamil_contracts_contract_kamil_contracts_waypay','waypay','contract_id',string="Claim", track_visibility="always")
    # waypay = fields.Many2one('kamil.contracts_waypay', track_visibility="always")
    #  For panalties
        
    # panalty_line_ids = fields.One2many('kamil.contracts.panalty.line', 'contract_id')
    panalty_ids = fields.One2many('kamil.contracts.panalty_lines','contract_id','contracts_penalty_lines', track_visibility="always")
    # penalty_lines = fields.Many2one('kamil.contracts.panalty_lines', track_visibility="always")


    #  For purchase contracts
    purchase_order_id = fields.Many2one('purchase.order', track_visibility="always")
    purchase_order_line_ids = fields.One2many('purchase.order.line', 'purchase_contract_id', track_visibility="always")
    # For medical contracts
    medical_services_ids = fields.One2many('kamil.contracts.services_price','medical_service_id', track_visibility="always")
    medical_ids = fields.One2many("medical.type","contract_medical_id")
    attachment_ids = fields.One2many('kamil.contracts.attachments','attachment_id',string='Attachment', track_visibility="always")
    
    error_end_date=fields.Char('Please select a date equal/or greater than the contract Start date')
    ext_rec_count = fields.Integer(compute='count_ext_func')
    amed_rec_count = fields.Integer(compute='count_amed_func')
    

    company_id = fields.Many2one('res.company', default= lambda self:self.env.user.company_id.id)
    study_mission_id = fields.Many2one("study.mission")
    

    _sql_constraints = [('contract_name_uniqe', 'unique (name,company_id)',
    'Sorry ! you can not create same name')]


    @api.multi
    def count_ext_func(self):
        # print ("Count of ext \n\n\n")
        # self.ensure_one()
        for rec in self:
           cnt = self.env['kamil.contracts.extensions'].search_count([('contract_id', '=', rec.id)])
           rec.ext_rec_count = cnt


    
    @api.multi
    def count_amed_func(self):
        self.ensure_one()
        for rec in self:
           cnt = self.env['kamil.contracts.contract'].search_count([('contract_id', '=', rec.id)])
           rec.amed_rec_count = cnt

    

    @api.multi
    @api.onchange('ntfs')
    def _onchange_ntfs(self):  
        employees = self.env['hr.employee'].search([('company_id','=',self.ntfs.id)])
        employees_list=[]
        for employee in employees:
            employees_list.append(employee.id)
        return {'domain':{'name_ntfs':[('id','in',employees_list)]}}
        
    
    @api.multi
    @api.onchange('second_party_ch')
    def _onchange_second_party_ch(self):  
        # print ("self.second_party_ch.id:::",self.second_party_ch.id,"\n\n\n")
        partners = self.env['res.partner'].search([('parent_id','=',self.second_party_ch.id)])
        partn_list=[]
        for partn in partners:
            partn_list.append(partn.id)
        return {'domain':{'second_party_name':[('id','in',partn_list)]}}


    
    @api.multi
    @api.onchange('purchase_order_id')
    def onchange_purchase_order(self):
        # purchase_order_line
        if self.purchase_order_id:
            sum_of_claims=0
            print("self.purchase_order_id::",self.purchase_order_id," total_amount",self.purchase_order_id.amount_total," \n\n")
            print("self.purchase_order_line_ids::",self.purchase_order_line_ids," \n\n\n")
            self.purchase_order_line_ids = False
            for order_line in self.purchase_order_id.order_line:
                order_line.purchase_contract_id = self.id 
                order_line.order_id = order_line.id 
                order_line.sum_purchase_order = order_line.product_qty * order_line.price_unit 
                

                sum_of_claims+=order_line.price_total
            self.purchase_order_id.amount_total=sum_of_claims
            self.contract_amount=sum_of_claims

            
        
    
                           
    
    # @api.model
    # def create(self, values):                
    #     values['barcode'] = self.env['ir.sequence'].get('kamil.contracts.barcode.req') 

    #     return super(kamilProjectContract, self).create(values)
  
   
class AddCapacity(models.Model):
    _inherit = 'res.partner'

    capacity = fields.Selection([('gm', 'General Manager'), ('cbd', 'Chairman of Board of Directors'),('delegated','Delegated'),('assignee','Assignee')],string="Capacity", default="gm")

    
class WayPay(models.Model):
    _name = 'kamil.contracts_waypay'
    _description='Payments Schedule'
    _inherit = ['mail.thread','mail.activity.mixin', 'portal.mixin']

    pay_date = fields.Date("Payment Date")
    execu_percentage = fields.Float('Complishment%',readonly=False)
    phase_name = fields.Char('Phase Name')
    contract_id = fields.Many2one('kamil.contracts.contract',ondelete='set null', string='Pay mode')
    # claim_contract_id= fields.Many2one('kamil.contracts.contract','Pay mode')
    status_sel = fields.Selection([('p', 'Partial Paid'),('f', 'Full Paid'), ('n', 'Not paid'),('i','in progress')],string="Payment status", default="n",readonly=True)
    pay_applied_amount = fields.Float('Applied Payment Amount',default=0)
    pay_paid_amount = fields.Float('Paid Amount')
    phase_amount = fields.Float('Phase Amount')
    claim_amount = fields.Float('Apply Amount', default=lambda self: self._compute_claim_balance())
    
    pay_balance = fields.Float('Payment Balance', compute='_compute_payment_balance')
    #Error Messages
    error_claim=fields.Char('Claim Amount cannot be > Your Balance for this Phase')
    error_claim='Claim Amount cannot be > Your Balance for this Phase'



    sum_of_claims=fields.Float()
    
    company_id = fields.Many2one('res.company', default= lambda self:self.env.user.company_id.id)
    # total_of_claim=fields.Float(compute='_compute_total_balance')

    _sql_constraints = [('contracts_waypay_name_uniqe', 'unique (phase_name,contract_id,company_id)',
     'Sorry ! you can not create same name')]
    
    _rec_name= 'phase_name'
    
    @api.one
    @api.constrains('claim_amount')
    def _claim_balance_constaint(self):
        for record in self:
            for rec in record:
                if rec.claim_amount > rec.pay_balance:
                    raise exceptions.ValidationError(self.error_claim)

        # total_of_phase_amount=0
        # for record in self.contract_id.payment_mode_ids:
        #     for rec in record:
        #         if rec:
        #             total_of_phase_amount+=rec.phase_amount
        #         rec.sum_of_claims=total_of_phase_amount
                
        
        # if rec.contract_id.contract_amount <total_of_phase_amount:
        #     raise exceptions.ValidationError(self.error_claim)
                    
    # @api.one
    # @api.constrains('claim_amount')
    # def _total_of_claims_constaint(self):
    #     total_of_phase_amount=0
    #     for record in self.contract_id.payment_mode_ids:
    #         for rec in record:
    #             if rec:
    #                 total_of_phase_amount+=rec.phase_amount
    #             rec.sum_of_claims=total_of_phase_amount
                
        
    #     if rec.contract_id.contract_amount <total_of_phase_amount:
    #         raise exceptions.ValidationError(self.error_claim)
    
    def _compute_claim_balance(self):
        for record in self:
            for rec in record:
                rec.claim_amount = rec.pay_balance
        

    @api.one
    def _compute_payment_balance(self):        
        for record in self:
            for rec in record:
                rec.pay_balance = rec.phase_amount - rec.pay_applied_amount    


class panalty_clause(models.Model):
    _name = 'kamil.contracts.panalty_lines'
    _rec_name= 'panalty_name'
    _inherit = ['mail.thread','mail.activity.mixin', 'portal.mixin']

    panalty_name = fields.Char('Panalty Name', track_visibility="always", required=True)
    panalty_amount = fields.Float('Panalty Amount', track_visibility="always")
    affect_date = fields.Date('Effective Date', track_visibility="always")
    is_active = fields.Boolean('Active', default=False, track_visibility="always")
    percentage = fields.Float('Percentage', track_visibility="always")
    frequency = fields.Selection([('onetime', 'One time'), ('foreachdayledday', 'For each delayed day')], string='Periodic', track_visibility="Frequency", required=True)
    percentage_of = fields.Selection([('c', 'Contract Amount'), ('p', 'Payment')], string='Percentage of', track_visibility="always")
    
    contract_id = fields.Many2one('kamil.contracts.contract',ondelete='set null', track_visibility="always")
    # contract_id2 = fields.Many2one('kamil.contracts.contract', track_visibility="always")
    panalty_basis = fields.Selection([('cont_per','Contract %'),('fixed_amont','Fixed amount')], track_visibility="always", required=True)

    # status_sel = fields.Selection([('p', 'Paid'), ('n', 'Not paid')],string="Penalty status", default='n', track_visibility="always")
    difference = fields.Integer(string='Date Difference', track_visibility="always")
    total_penalty = fields.Float(string='Total', track_visibility="always")
    opened = fields.Boolean(string='Opened',default=False, track_visibility="always")
    contract_state= fields.Char( default=lambda self: self.contract_id.state)
    company_id = fields.Many2one('res.company', default= lambda self:self.env.user.company_id.id)
    

    _sql_constraints = [('name_uniqe', 'unique (panalty_name,contract_id)','Sorry ! you can not create same name')]

    panalty_basis
    @api.multi
    @api.onchange('panalty_basis')
    def onchnage_panalty_basis(self):
        if self.panalty_basis=='cont_per':
            if self.contract_id.contract_amount and self.contract_id.contract_amount>0:
                # self.percentage = self.contract_id.contract_amount / self.percentage *100
                self.panalty_amount = self.contract_id.contract_amount * self.percentage /100
        
        elif self.panalty_basis=='fixed_amont':
            self.percentage = (self.panalty_amount / self.contract_id.contract_amount) *100    

    
    @api.multi
    def write(self, values):
        if 'panalty_amount' in values:
            if values['panalty_amount'] > 0.0:
                values['percentage'] = 0
        if 'percentage' in values:
            if values['percentage'] >0:
                values['panalty_amount'] = 0.0
        return super(panalty_clause, self).write(values)



# class panaltyLine(models.Model):
#     _name = 'kamil.contracts.panalty.line'
#     _rec_name= 'panalty_name'

#     panalty_name = fields.Char('Panalty Name', track_visibility="always")
#     panalty_amount = fields.Float('Panalty Amount', track_visibility="always")
#     affect_date = fields.Date('Panalty Date',readonly=1, track_visibility="always")
#     is_active = fields.Boolean('Active', default=False, track_visibility="always")
#     percentage = fields.Integer('Percentage%', track_visibility="always")
#     percentage_of = Selection([('c', 'Contract Amount'), ('p', 'Payment')], string='Percentage of', track_visibility="always")
#     contract_id = fields.Many2one('kamil.contracts.contract', track_visibility="always")
#     contract_id2 = fields.Many2one('kamil.contracts.contract', track_visibility="always")



class PurchaseRequisitionLine(models.Model):
    _inherit = 'purchase.order.line'

    purchase_contract_id = fields.Many2one('kamil.contracts.contract', track_visibility="always", ondelete='restrict')
    sum_purchase_order= fields.Float()
    # product_qty = fields.Char(required=True, track_visibility="always")
    # price_unit= fields.Char(required=True, track_visibility="always")
    # product_id= fields.Char(required=True, track_visibility="always")
    company_id = fields.Many2one('res.company', default= lambda self:self.env.user.company_id.id)
class kamilContractServicesPrice(models.Model):
    _name = 'kamil.contracts.services_price'
    _rec_name = 'e_service_name'
    _inherit = ['mail.thread','mail.activity.mixin', 'portal.mixin']
    
    e_service_name = fields.Char(string='English Name', required=True, track_visibility="always")
    a_service_name = fields.Char(string='Arabic Name', required=True, track_visibility="always")
    service_code = fields.Char(string='Service Code', required=True, track_visibility="always")
    service_scode = fields.Char(required=True)

    service_price_for_public = fields.Integer('Price for Public', required=True, track_visibility="always")
    service_price_for_employees = fields.Integer('Price for Employees', required=True, track_visibility="always")

    medical_service_id = fields.Many2one('kamil.contracts.contract', track_visibility="always")
    _sql_constraints = [('medical_service_name_uniqe', 'unique (e_service_name)','Sorry ! you can not create same name')]

class MedicalGroup(models.Model):
    _name = "medical.group"
    _rec_name = "MainGroup"
    MainGroup = fields.Char(required=True)
    SubGroup = fields.Many2one("medical.group")

class FormType(models.Model):
    _name = "form.type"
    name = fields.Char(required=True)

class MedicalType(models.Model):
    _name = "medical.type"
    name = fields.Char(required=True)
    MainGroup_id = fields.Many2one("medical.group",required=True)
    SubGroup_id = fields.Many2one("medical.group",required=True)
    form_id = fields.Many2one("form.type",required=True)
    strength = fields.Char(required=True)
    contract_medical_id = fields.Many2one("kamil.contracts.contract")
    
    price_for_public = fields.Float('Public Price')
    price_for_private = fields.Float('Private Price')

    # _sql_constraints = [('medical_uniqe_name', 'unique (name,MainGroup_id,SubGroup_id,form_id,strength)','Sorry ! you can not create same name')]

    @api.multi
    @api.onchange('MainGroup_id')
    def onchange_MainGroup_id(self):
        if self.MainGroup_id:
            return{
            'domain':{
                'SubGroup_id':[('SubGroup','=',self.MainGroup_id.id)]
            }
        }

class kamilContractAttachments(models.Model):
    _name = 'kamil.contracts.attachments'
    _rec_name ='attachment_name'
    _inherit = ['mail.thread','mail.activity.mixin', 'portal.mixin']    
    
    
    attachment_name = fields.Char(string='Attachment Name', required=True, track_visibility="always")
    attachment_id =fields.Many2one('kamil.contracts.contract', track_visibility="always", ondelete='restrict')
    attachment = fields.Binary( track_visibility="always")
    company_id = fields.Many2one('res.company', default= lambda self:self.env.user.company_id.id)
    _sql_constraints = [('attachment_uniqe_name', 'unique (attachment_name,attachment_id)','Sorry ! you can not create same name')]